from flask import Blueprint, render_template, flash, request, redirect, url_for, send_file
from docx import Document
from io import BytesIO, StringIO

views = Blueprint('views', __name__)
# global classes
doc = Document()


def create_timetable(days, times, classes):
    doc.add_heading('College Class Timetable', level=1)

    table = doc.add_table(rows=len(days) + 1, cols=len(times) + 1)
    table.style = 'Table Grid'

    header_row = table.rows[0]
    header_row.cells[0].text = 'Day / Time'
    for i, time in enumerate(times):
        header_row.cells[i+1].text = time

    for j, day in enumerate(days):
        row = table.rows[j+1].cells
        row[0].text = day
        for i, time in enumerate(times):
            subject = classes.get(day, {}).get(time, "")
            row[i+1].text = subject


def add_course_info_table(course_data):
    doc.add_heading('Course Information', level=1)

    table = doc.add_table(rows=len(course_data) + 1, cols=5)
    table.style = 'Table Grid'

    header_row = table.rows[0]
    header_row.cells[0].text = 'Course Code'
    header_row.cells[1].text = 'Course Name'
    header_row.cells[2].text = 'No. of Hours'
    header_row.cells[3].text = 'Faculty Name'
    header_row.cells[4].text = 'Department'

    for i, (code, name, hours, faculty, department) in enumerate(course_data, start=1):
        row = table.rows[i]
        row.cells[0].text = code
        row.cells[1].text = name
        row.cells[2].text = hours
        row.cells[3].text = faculty
        row.cells[4].text = department


@views.route('/', methods=['POST', 'GET'])
def home():
    if request.method == 'POST':
        times = request.form.get('slots').split(',')
        return redirect(url_for('views.time_table', timesz=times))
    return render_template("home.html")


@views.route('timetable/<string:timesz>', methods=['GET', 'POST'])
def time_table(timesz):
    if request.method == 'POST':
        classes = {}
        str1 = ""
        for i in range(1, len(timesz) - 1):
            if ord(timesz[i]) == ord('"') or ord(timesz[i]) == ord("'"):
                pass
            else:
                str1 += timesz[i]
        sp = str1.split(',')
        tim = [time.strip() for time in sp]
        days = ['Monday', 'Tuesday', 'Wednesday', 'Thursday', 'Friday']
        for day in days:
            classes[day] = {}
            for time in tim:
                classes[day][time] = request.form.get(f"{day}_{time}")
        create_timetable(days, tim, classes)
        return redirect(url_for('views.total_course'))
    days = ['Monday', 'Tuesday', 'Wednesday', 'Thursday', 'Friday']
    str1 = ""
    for i in range(1, len(timesz)-1):
        if ord(timesz[i]) == ord('"') or ord(timesz[i]) == ord("'"):
            pass
        else:
            str1 += timesz[i]
    sp = str1.split(',')
    tim = [time.strip() for time in sp]
    return render_template("timetable.html", days=days, times=tim)


@views.route('total_course', methods=['GET', 'POST'])
def total_course():
    if request.method == 'POST':
        course_count = request.form.get("course_count")
        return redirect(url_for("views.course_info", num_courses=course_count))
    return render_template("total_course.html")


@views.route('course_info/<int:num_courses>', methods=['GET', 'POST'])
def course_info(num_courses):
    if request.method == 'POST':
        course_data = []
        for i in range(num_courses):
            code = request.form.get(f"code_{i+1}")
            name = request.form.get(f"name_{i+1}")
            hours = request.form.get(f"hours_{i+1}")
            faculty = request.form.get(f"faculty_{i+1}")
            department = request.form.get(f"department_{i+1}")
            course_data.append((code, name, hours, faculty, department))
        add_course_info_table(course_data)
        doc_io = BytesIO()
        doc.save(doc_io)
        doc_io.seek(0)
        filename = 'college_timetable_and_courses.docx'
        return send_file(doc_io, as_attachment=True, download_name=filename, mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    return render_template("course_info.html", num_courses=num_courses)
